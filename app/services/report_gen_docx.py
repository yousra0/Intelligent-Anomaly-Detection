"""
app/services/report_gen_docx.py
Génère le rapport Word en remplissant le template exemple_rapport.docx.
Placeholders {{...}}, graphiques matplotlib, tableau et sections dynamiques.
"""
from __future__ import annotations

import io
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TEMPLATE_PATH = Path(__file__).parent.parent.parent / "exemple_rapport.docx"

# ─── Couleurs PwC ─────────────────────────────────────────────────────────────
_C_ORANGE = "#D04A02"
_C_DARK   = "#293854"
_C_RED    = "#C00000"
_C_WARN   = "#E66E00"
_C_GREEN  = "#008246"
_C_LGRAY  = "#F8F8F8"
_C_MGRAY  = "#B0B0B0"

RGB_ORANGE = RGBColor(208, 74,  2)
RGB_DARK   = RGBColor(41,  56,  84)
RGB_RED    = RGBColor(192, 0,   0)
RGB_WARN   = RGBColor(230, 110, 0)
RGB_GREEN  = RGBColor(0,   130, 70)
RGB_WHITE  = RGBColor(255, 255, 255)
RGB_GRAY   = RGBColor(80,  80,  80)

FEATURE_FR = {
    "balance_diff_orig": "Ecart de solde (emetteur)",
    "log_amount":        "Montant (echelle log)",
    "amount":            "Montant de la transaction",
    "hour":              "Heure de la transaction",
    "dest_zero_balance": "Solde nul chez le destinataire",
    "type_CASH_OUT":     "Type : Retrait especes",
    "type_TRANSFER":     "Type : Virement bancaire",
    "type_CASH_IN":      "Type : Depot especes",
    "type_PAYMENT":      "Type : Paiement",
    "type_DEBIT":        "Type : Debit automatique",
}

FEATURE_EXPLAIN = {
    "balance_diff_orig": (
        "Le solde du compte emetteur a change d'une facon qui ne correspond pas "
        "aux operations declarees. Indicateur tres fort de fraude potentielle."
    ),
    "log_amount": (
        "Le montant est inhabituellement eleve par rapport aux transactions similaires."
    ),
    "hour": (
        "Transaction effectuee a une heure atypique (nuit / tres tot le matin), "
        "periode connue pour concentrer les activites frauduleuses."
    ),
    "dest_zero_balance": (
        "Compte destinataire avec solde nul apres reception : caracteristique "
        "des comptes 'mule' utilises pour blanchiment."
    ),
    "type_CASH_OUT": "Les retraits especes sont la methode la plus frequente dans les fraudes.",
    "type_TRANSFER": "Virements de grande valeur vers comptes inconnus : haut risque.",
}

RISK_LABEL = {
    "CRITIQUE": "A investiguer immediatement",
    "ELEVE":    "A surveiller prioritairement",
    "FAIBLE":   "Aucun probleme detecte",
}

GLOSSAIRE = [
    ("Intelligence artificielle (IA)",
     "Programme informatique capable d'apprendre a partir de donnees historiques "
     "pour detecter des comportements anormaux, sans etre explicitement programme."),
    ("XGBoost",
     "Algorithme d'apprentissage automatique qui analyse de nombreux indicateurs "
     "simultanement pour attribuer un score de risque (0 = normal, 1 = fraude certaine)."),
    ("AutoEncoder",
     "Reseau de neurones qui apprend a 'reconstruire' des transactions normales. "
     "Quand il echoue, cela signale une anomalie."),
    ("SHAP - Facteurs declencheurs",
     "Technique qui explique pourquoi l'IA a signale une transaction en identifiant "
     "les indicateurs les plus contributifs a l'alerte."),
    ("Indice de risque (0 a 1)",
     "Score attribue par l'IA. Plus proche de 1 = plus suspect. "
     "Un seuil calibre sur donnees historiques determine l'alerte."),
    ("Solde nul destinataire",
     "Compte recevant l'argent avec solde nul avant la transaction. "
     "Caracteristique souvent associee aux comptes utilises pour blanchiment."),
    ("Ecart de solde",
     "Difference entre le montant debite du compte source et la variation de solde. "
     "Un ecart important peut indiquer une manipulation des donnees."),
]


# ─── Helpers matplotlib ───────────────────────────────────────────────────────

def _to_buf(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    buf.seek(0)
    plt.close(fig)
    return buf


def _chart_donut(n_tx: int, n_fraud: int) -> io.BytesIO:
    n_ok   = max(n_tx - n_fraud, 0)
    sizes  = [n_ok, n_fraud] if n_fraud > 0 else [1, 0]
    labels = [f"Normales\n{n_ok:,}", f"Suspectes\n{n_fraud:,}"]
    fig, ax = plt.subplots(figsize=(4.0, 3.2))
    fig.patch.set_facecolor("white")
    _, _, autotexts = ax.pie(
        sizes, labels=labels, colors=[_C_DARK, _C_ORANGE],
        autopct="%1.1f%%", startangle=90, pctdistance=0.78,
        wedgeprops=dict(width=0.52, edgecolor="white", linewidth=2),
        textprops=dict(fontsize=8),
    )
    for at in autotexts:
        at.set_color("white")
        at.set_fontweight("bold")
        at.set_fontsize(8)
    ax.set_title("Repartition des transactions", fontsize=9, fontweight="bold",
                 color=_C_DARK, pad=8)
    plt.tight_layout(pad=0.3)
    return _to_buf(fig)


def _chart_risk_bars(transactions: list[dict]) -> io.BytesIO:
    counts = Counter(t.get("risk_level", "FAIBLE") for t in transactions)
    cats   = ["FAIBLE", "ELEVE", "CRITIQUE"]
    labels = ["Normale", "A surveiller", "Critique"]
    vals   = [counts.get(c, 0) for c in cats]
    hues   = [_C_GREEN, _C_WARN, _C_RED]
    max_v  = max(vals) if max(vals) > 0 else 1

    fig, ax = plt.subplots(figsize=(4.0, 3.2))
    fig.patch.set_facecolor("white")
    ax.set_facecolor(_C_LGRAY)
    bars = ax.bar(labels, vals, color=hues, edgecolor="white", linewidth=1.5, width=0.5)
    for bar, v in zip(bars, vals):
        if v > 0:
            ax.text(bar.get_x() + bar.get_width() / 2,
                    bar.get_height() + max_v * 0.03,
                    str(v), ha="center", va="bottom",
                    fontsize=9, fontweight="bold", color=_C_DARK)
    ax.set_title("Niveau de risque detecte", fontsize=9, fontweight="bold",
                 color=_C_DARK, pad=8)
    ax.set_ylabel("Transactions", fontsize=7, color="#666666")
    ax.set_ylim(0, max_v * 1.25 + 1)
    ax.tick_params(labelsize=7.5)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    plt.tight_layout(pad=0.3)
    return _to_buf(fig)


def _chart_amount_type(transactions: list[dict]) -> io.BytesIO:
    sums: dict[str, float] = defaultdict(float)
    for t in transactions:
        if t.get("risk_level") in ("CRITIQUE", "ELEVE") or t.get("is_fraud_predicted"):
            sums[t.get("type", "Autre")] += t.get("amount", 0)
    if not sums:
        for t in transactions:
            sums[t.get("type", "Autre")] += t.get("amount", 0)
    items = sorted(sums.items(), key=lambda x: -x[1])[:6] or [("Aucune donnee", 0)]
    types = [i[0] for i in items]
    amts  = [i[1] / 1_000 for i in items]
    pal   = [_C_ORANGE, _C_DARK, _C_RED, _C_WARN, _C_GREEN, _C_MGRAY]
    max_a = max(amts) if max(amts) > 0 else 1

    fig, ax = plt.subplots(figsize=(6.5, 3.2))
    fig.patch.set_facecolor("white")
    ax.set_facecolor(_C_LGRAY)
    ax.barh(types[::-1], amts[::-1], color=pal[:len(types)][::-1],
            edgecolor="white", linewidth=1.2)
    for bar, v in zip(ax.patches, amts[::-1]):
        ax.text(bar.get_width() + max_a * 0.02,
                bar.get_y() + bar.get_height() / 2,
                f"{v:,.0f}k TND", va="center", ha="left",
                fontsize=7.5, color=_C_DARK)
    ax.set_title("Exposition financiere par type d'operation (kTND)",
                 fontsize=9, fontweight="bold", color=_C_DARK, pad=8)
    ax.set_xlabel("Milliers TND", fontsize=7, color="#666666")
    ax.set_xlim(0, max_a * 1.4)
    ax.tick_params(labelsize=7.5)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    plt.tight_layout(pad=0.3)
    return _to_buf(fig)


def _chart_score_dist(transactions: list[dict], threshold: float = 0.355) -> io.BytesIO:
    """Histogram of AI risk scores — the most explanatory chart for non-technical auditors."""
    scores = [float(t.get("xgb_score", 0)) for t in transactions] or [0.0]
    normal     = [s for s in scores if s <  threshold]
    suspicious = [s for s in scores if s >= threshold]
    bins = [i / 20 for i in range(21)]

    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    fig.patch.set_facecolor("white")
    ax.set_facecolor(_C_LGRAY)
    ax.hist(normal,     bins=bins, color=_C_DARK,   alpha=0.85,
            label=f"Normales ({len(normal):,})", edgecolor="white", linewidth=0.5)
    ax.hist(suspicious, bins=bins, color=_C_ORANGE, alpha=0.90,
            label=f"Suspectes ({len(suspicious):,})", edgecolor="white", linewidth=0.5)
    ax.axvline(x=threshold, color=_C_RED, linestyle="--", linewidth=1.8,
               label=f"Seuil IA = {threshold:.3f}")
    ax.set_title(
        "Distribution des indices de risque IA\n"
        "(0 = transaction normale  |  1 = fraude quasi-certaine)",
        fontsize=9, fontweight="bold", color=_C_DARK, pad=8,
    )
    ax.set_xlabel("Indice de risque (0 -> 1)", fontsize=8, color="#666666")
    ax.set_ylabel("Nombre de transactions", fontsize=8, color="#666666")
    ax.legend(fontsize=8, framealpha=0.7)
    ax.tick_params(labelsize=8)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    plt.tight_layout(pad=0.4)
    return _to_buf(fig)


def _chart_shap(explain_results: list[dict]) -> io.BytesIO | None:
    totals: dict[str, float] = defaultdict(float)
    cnts:   dict[str, int]   = defaultdict(int)
    for e in explain_results:
        for f, v in e.get("shap_values", {}).items():
            totals[f] += abs(float(v))
            cnts[f]   += 1
    if not totals:
        return None
    mean_abs = {f: totals[f] / cnts[f] for f in totals}
    items  = sorted(mean_abs.items(), key=lambda x: -x[1])[:8]
    labels = [FEATURE_FR.get(f, f) for f, _ in items]
    vals   = [v for _, v in items]
    thresh = vals[0] * 0.35 if vals else 0
    hues   = [_C_ORANGE if v >= thresh else _C_DARK for v in vals]

    fig, ax = plt.subplots(figsize=(7.0, 3.2))
    fig.patch.set_facecolor("white")
    ax.set_facecolor(_C_LGRAY)
    ax.barh(labels[::-1], vals[::-1], color=hues[::-1], edgecolor="white", linewidth=1)
    ax.set_title("Facteurs les plus declencheurs d'alertes (importance SHAP)",
                 fontsize=9, fontweight="bold", color=_C_DARK, pad=8)
    ax.set_xlabel("Importance moyenne", fontsize=7, color="#666666")
    ax.tick_params(labelsize=7.5)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    plt.tight_layout(pad=0.3)
    return _to_buf(fig)


def _fmt_amount(v: float) -> str:
    if v >= 1_000_000:
        return f"{v / 1_000_000:.2f}M"
    if v >= 1_000:
        return f"{v / 1_000:.0f}k"
    return f"{v:,.0f}"


# ─── Helpers docx ─────────────────────────────────────────────────────────────

def _para_text(para) -> str:
    return "".join(r.text or "" for r in para.runs)


def _replace_in_para(para, placeholder: str, new_text: str) -> bool:
    """Replace placeholder handling run-split (Word splits {{...}} across runs)."""
    full = _para_text(para)
    if placeholder not in full:
        return False
    new_full = full.replace(placeholder, new_text)
    if para.runs:
        para.runs[0].text = new_full
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_full)
    return True


_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _replace_all(doc: Document, mapping: dict[str, str]) -> None:
    """Replace placeholders in every <w:p> across the entire document body.

    Word 365 splits {{key}} across multiple <w:r> elements with <w:proofErr>
    spell-check markers between them.  Concatenating all run texts per paragraph,
    replacing, then writing back to the first <w:t> handles this for every
    context: VML v:textbox shapes, DrawingML anchors, and plain body paragraphs.
    """
    for p in doc.element.body.iter(f"{{{_NS_W}}}p"):
        r_elems = p.findall(f"{{{_NS_W}}}r")
        valid_t = [r.find(f"{{{_NS_W}}}t") for r in r_elems]
        valid_t = [t for t in valid_t if t is not None]
        if not valid_t:
            continue

        full = "".join(t.text or "" for t in valid_t)
        if "{{" not in full:
            continue

        new_full = full
        for ph, val in mapping.items():
            if ph in new_full:
                new_full = new_full.replace(ph, val)

        if new_full != full:
            valid_t[0].text = new_full
            for t in valid_t[1:]:
                t.text = ""


def _find_para(doc: Document, substr: str):
    for para in doc.paragraphs:
        if substr in para.text:
            return para
    return None


def _fill_para_with_chart(para, img_buf: io.BytesIO,
                          doc: Document, width: float = 5.5) -> None:
    """Clear a paragraph's text and insert a chart image inline.
    Works on a pre-saved paragraph reference so it is immune to the
    cleared-placeholder problem (placeholder already gone after _replace_all)."""
    for r in para._element.findall(f"{{{_NS_W}}}r"):
        t = r.find(f"{{{_NS_W}}}t")
        if t is not None:
            t.text = ""
    for pe in list(para._element.findall(f"{{{_NS_W}}}proofErr")):
        para._element.remove(pe)
    run = para.add_run()
    run.add_picture(img_buf, width=Inches(width))


def _ensure_page_break_before(para) -> None:
    """Add w:pageBreakBefore to paragraph properties so this paragraph
    always starts on a new page regardless of prior content length."""
    pPr = para._element.find(f"{{{_NS_W}}}pPr")
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._element.insert(0, pPr)
    if pPr.find(f"{{{_NS_W}}}pageBreakBefore") is None:
        pbk = OxmlElement("w:pageBreakBefore")
        pPr.append(pbk)


def _clear_empty_paras_after(doc: Document, ref_para, count: int = 9) -> None:
    paras = doc.paragraphs
    try:
        idx = paras.index(ref_para)
    except ValueError:
        return
    to_remove = []
    for para in paras[idx + 1: idx + 1 + count]:
        if not para.text.strip():
            to_remove.append(para._element)
        else:
            break
    for elem in to_remove:
        elem.getparent().remove(elem)


def _set_cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_color)
    shd.set(qn("w:color"), hex_color)
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)


def _cell_text(cell, text: str, bold: bool = False,
               color: RGBColor = None, size: int = 11,
               align: str = "center") -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = {"center": 1, "left": 0, "right": 2}.get(align, 0)
    run  = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _no_borders_table(table) -> None:
    """Remove all visible borders from a table."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    bdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        bdr.append(el)
    tblPr.append(bdr)


def _insert_charts_grid(doc: Document, anchor_elem,
                         charts: list[tuple[io.BytesIO, str]],
                         width_each: float = 2.9) -> None:
    """Insert charts 2-per-row in a borderless table right after anchor_elem.
    Each entry in charts is (image_buf, caption_label)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    rows_data = [charts[i:i + 2] for i in range(0, len(charts), 2)]
    table = doc.add_table(rows=len(rows_data), cols=2)
    _no_borders_table(table)

    for ri, pair in enumerate(rows_data):
        for ci, (buf, label) in enumerate(pair):
            buf.seek(0)
            cell = table.cell(ri, ci)
            # small caption above
            cap = cell.paragraphs[0]
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap.add_run(label)
            cap_run.bold = True
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGB_DARK
            # chart image
            img_p = cell.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.add_run().add_picture(buf, width=Inches(width_each))

    anchor_elem.addnext(table._tbl)


def _clear_paras_between(doc: Document, start_para, end_para) -> None:
    """Remove every paragraph strictly between start_para and end_para."""
    paras = doc.paragraphs
    try:
        si = paras.index(start_para) + 1
        ei = paras.index(end_para)
    except ValueError:
        return
    for para in paras[si:ei]:
        para._element.getparent().remove(para._element)


def _add_para_after(ref_elem, text: str = "", bold: bool = False,
                    italic: bool = False, size: int = 11,
                    color: RGBColor = None, indent: float = 0.0) -> OxmlElement:
    """Insert a new <w:p> element right after ref_elem. Returns the new element."""
    new_p   = OxmlElement("w:p")
    new_pPr = OxmlElement("w:pPr")

    if indent > 0:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(int(indent * 914400 / 2.54 / 10)))  # cm to twips
        new_pPr.append(ind)

    new_p.append(new_pPr)

    if text:
        new_r   = OxmlElement("w:r")
        new_rPr = OxmlElement("w:rPr")

        if bold:
            b = OxmlElement("w:b")
            new_rPr.append(b)
        if italic:
            i = OxmlElement("w:i")
            new_rPr.append(i)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size * 2))
        new_rPr.append(sz)
        if color:
            clr = OxmlElement("w:color")
            clr.set(qn("w:val"), str(color))   # RGBColor.__str__ -> "RRGGBB"
            new_rPr.append(clr)

        new_r.append(new_rPr)
        new_t = OxmlElement("w:t")
        new_t.text = text
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_r.append(new_t)
        new_p.append(new_r)

    ref_elem.addnext(new_p)
    return new_p


def _add_table_after_para(doc: Document, ref_para,
                           headers: list[str], rows: list[list[str]],
                           col_widths: list[float] | None = None) -> None:
    """Add a styled table immediately after ref_para."""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "TableGrid"

    # Header row
    hdr_row = table.rows[0]
    for ci, (h, cell) in enumerate(zip(headers, hdr_row.cells)):
        _set_cell_bg(cell, "293854")
        _cell_text(cell, h, bold=True, color=RGB_WHITE)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = "F8F8F8" if ri % 2 == 0 else "FFFFFF"
        for ci, (val, cell) in enumerate(zip(row_data, row.cells)):
            _set_cell_bg(cell, bg)
            if ci == n_cols - 1:
                col = RGB_RED if "ALERTE" in val else (RGB_WARN if "ATTENTION" in val else RGB_GREEN)
                _cell_text(cell, val, bold=True, color=col)
            else:
                _cell_text(cell, val)

    # Set column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)

    # Move table to just after ref_para
    ref_para._element.addnext(table._tbl)


# ─── Générateur principal ─────────────────────────────────────────────────────

def generate_pwc_docx_report(
    predict_result: dict,
    explain_results: list[dict] | None = None,
) -> bytes:
    now          = datetime.now()
    date_str     = now.strftime("%d/%m/%Y a %H:%M")

    n_tx         = predict_result.get("n_transactions", 0)
    n_fraud      = predict_result.get("n_fraud", 0)
    rate         = predict_result.get("fraud_rate_pct", 0.0)
    amount_risk  = predict_result.get("amount_at_risk", 0.0)
    threshold    = predict_result.get("threshold", 0.355)
    transactions = predict_result.get("transactions", [])

    risk_counts  = Counter(t.get("risk_level", "FAIBLE") for t in transactions)
    n_critique   = risk_counts.get("CRITIQUE", 0)
    n_eleve      = risk_counts.get("ELEVE",    0)
    n_faible     = risk_counts.get("FAIBLE",   0)
    total        = max(n_tx, 1)

    if rate > 5:
        global_risk = "CRITIQUE"
    elif rate > 1:
        global_risk = "ELEVE"
    else:
        global_risk = "FAIBLE"

    # ── Textes descriptifs ────────────────────────────────────────────────────
    resume_text = (
        "Notre systeme d'intelligence artificielle a analyse l'ensemble des transactions "
        "soumises et identifie les operations dont les caracteristiques s'ecartent "
        "significativement des comportements normaux. Les resultats sont presentes en "
        "trois categories : transactions normales, a surveiller et critiques. "
        f"Sur {n_tx:,} transactions analysees, {n_fraud:,} ont ete signalees comme suspectes, "
        f"representant {rate:.2f}% du portefeuille et un montant total de "
        f"{amount_risk:,.2f} TND exposes."
    )

    analyse_text = (
        "Les graphiques ci-dessous permettent de visualiser rapidement la distribution "
        "des transactions, les niveaux de risque detectes et l'exposition financiere "
        "par type d'operation. Chaque graphique est accompagne de son interpretation "
        "pour faciliter la prise de decision de l'auditeur."
    )

    transaction_text = (
        "Le tableau ci-dessous presente les 10 transactions les plus suspectes, "
        "classees par indice de risque decroissant. Les transactions marquees 'ALERTE' "
        "necessitent une verification immediate. L'indice IA varie de 0 (transaction "
        "normale) a 1 (fraude quasi-certaine)."
    )

    # ── Remplacement des placeholders ────────────────────────────────────────
    doc = Document(str(TEMPLATE_PATH))

    # ── Références AVANT tout remplacement ───────────────────────────────────
    para_analyse   = _find_para(doc, "{{analyse_text}}")
    para_tx_text   = _find_para(doc, "{{transaction_text}}")
    para_section02 = _find_para(doc, "02 Analyse")
    para_section03 = _find_para(doc, "03 Transactions")
    para_section04 = _find_para(doc, "04 Fiche")
    para_section05 = _find_para(doc, "05 Recommandations")
    para_section06 = _find_para(doc, "06 Glossaire")
    para_rec       = _find_para(doc, "{{recommendations_text}}")
    para_glo       = _find_para(doc, "{{glossaire_")

    # ── Sauts de page entre sections (contenu = 1 page par section) ──────────
    for sec_para in (para_section02, para_section03, para_section04,
                     para_section05, para_section06):
        if sec_para:
            _ensure_page_break_before(sec_para)

    mapping = {
        "{{date}}":                  date_str,
        "{{resume_text}}":           resume_text,
        "{{analyse_text}}":          analyse_text,
        "{{transaction_text}}":      transaction_text,
        # placeholders texte dynamiques : vidés -> contenu injecté après
        "{{recommendations_text}}":  "",
        "{{glossaire_ text }}":      "",
        # placeholders graphiques : les paragraphes entre analyse et section03
        # seront supprimés par _clear_paras_between, donc pas besoin de les vider ici
        "{{reparation_transactions_graphes}}":  "",
        "{{niveau_risque_graphes}}":            "",
        "{{exposition_financière_graphes}}":    "",
        # KPI numériques (boîtes VML)
        "{{taux_de_fraudes}}":       f"{rate:.2f}%",
        "{{anomalies_détéctées}}":   f"{n_fraud:,}",
        "{{transactions_analysées}}": f"{n_tx:,}",
        "{{montant_exposé}}":        _fmt_amount(amount_risk),
        # jauge de risque
        "{{pourcentage}}":           f"{rate:.2f}%",
        "{{remarque}}":              RISK_LABEL[global_risk].upper(),
        "{{pourcentage_faible}}":    f"{n_faible / total * 100:.0f}%",
        "{{pourcentage_elevé}}":     f"{n_eleve  / total * 100:.0f}%",
        "{{pourcentage_critique}}":  f"{n_critique / total * 100:.0f}%",
    }
    _replace_all(doc, mapping)

    # ── Graphiques ─────────────────────────────────────────────────────────────
    buf_donut  = _chart_donut(n_tx, n_fraud)
    buf_bars   = _chart_risk_bars(transactions)
    buf_types  = _chart_amount_type(transactions)
    buf_scores = _chart_score_dist(transactions, threshold)

    charts: list[tuple[io.BytesIO, str]] = [
        (buf_donut,  "Repartition des transactions"),
        (buf_bars,   "Niveaux de risque detectes"),
        (buf_types,  "Exposition financiere par type (kTND)"),
        (buf_scores, "Distribution des indices de risque IA"),
    ]
    if explain_results:
        buf_shap = _chart_shap(explain_results)
        if buf_shap:
            charts.append((buf_shap, "Facteurs declencheurs (SHAP global)"))

    # Supprimer tous les paragraphes entre analyse_text et section 03
    # (labels, placeholders charts vides, espacements) puis insérer la grille
    if para_analyse and para_section03:
        _clear_paras_between(doc, para_analyse, para_section03)
        _insert_charts_grid(doc, para_analyse._element, charts, width_each=2.9)

    # ── Tableau transactions ──────────────────────────────────────────────────
    if para_tx_text and transactions:
        top10 = sorted(transactions, key=lambda t: t.get("xgb_score", 0), reverse=True)[:10]
        headers = ["#", "Identifiant", "Type", "Montant (TND)", "Indice IA", "Statut"]
        widths  = [0.25, 1.10, 0.85, 1.10, 0.95, 1.15]
        rows = []
        for rank, tx in enumerate(top10, 1):
            risk = tx.get("risk_level", "FAIBLE")
            icon = {"CRITIQUE": "ALERTE", "ELEVE": "ATTENTION", "FAIBLE": "OK"}[risk]
            rows.append([
                str(rank),
                str(tx.get("tx_id", ""))[:14],
                str(tx.get("type", "-")),
                f"{tx.get('amount', 0):,.0f}",
                f"{tx.get('xgb_score', 0):.3f}",
                icon,
            ])
        _add_table_after_para(doc, para_tx_text, headers, rows, col_widths=widths)
        _clear_empty_paras_after(doc, para_tx_text, count=9)

    # ── Fiches critiques (tableau compact, une seule page) ───────────────────
    explains_map: dict[str, dict] = {}
    if explain_results:
        for e in explain_results:
            key = str(e.get("tx_id", e.get("id", "")))
            explains_map[key] = e

    critique_txs = [t for t in transactions if t.get("risk_level") == "CRITIQUE"]

    if para_section04:
        _clear_empty_paras_after(doc, para_section04, count=15)

        if not critique_txs:
            _add_para_after(
                para_section04._element,
                "Aucune transaction critique detectee lors de cette analyse.",
                italic=True, size=11, color=RGB_GRAY,
            )
        else:
            hdrs   = ["#", "ID Transaction", "Type", "Montant (TND)",
                      "Indice IA", "Facteur principal", "Statut"]
            widths = [0.25, 1.10, 0.80, 1.10, 0.75, 1.90, 0.60]
            rows   = []
            for rank, tx in enumerate(critique_txs[:14], 1):
                tx_id  = str(tx.get("tx_id", f"TX-{rank}"))[:14]
                expl   = explains_map.get(tx_id, {})
                shap_v = expl.get("shap_values", {})
                llm    = expl.get("llm", expl)
                if shap_v:
                    top_f   = max(shap_v.items(), key=lambda x: abs(float(x[1])))[0]
                    facteur = FEATURE_FR.get(top_f, top_f)[:28]
                elif llm.get("raisons"):
                    facteur = str(llm["raisons"][0])[:28]
                else:
                    facteur = "-"
                rows.append([
                    str(rank), tx_id,
                    str(tx.get("type", "-")),
                    f"{tx.get('amount', 0):,.0f}",
                    f"{float(tx.get('xgb_score', 0)):.3f}",
                    facteur,
                    "ALERTE",
                ])
            _add_table_after_para(doc, para_section04, hdrs, rows, col_widths=widths)

    # ── Recommandations ───────────────────────────────────────────────────────
    if para_rec:
        recs = []
        if n_critique > 0:
            recs.append(
                f"Investiguer immediatement les {n_critique} transaction(s) marquee(s) 'CRITIQUE'. "
                "Verifiez la legitimite des beneficiaires et les pieces justificatives associees."
            )
        if n_eleve > 0:
            recs.append(
                f"Planifier une revue approfondie des {n_eleve} transaction(s) 'A SURVEILLER' "
                "dans les 48 heures suivantes."
            )
        if amount_risk > 0:
            recs.append(
                f"Le montant total expose est de {amount_risk:,.2f} TND. "
                "Evaluez l'impact financier potentiel avec le responsable financier."
            )
        recs += [
            "Documenter toutes vos verifications dans le dossier d'audit permanent.",
            "En cas de fraude confirmee, declencher la procedure d'alerte interne.",
            "Conserver ce rapport joint aux preuves collectees dans le dossier d'audit.",
        ]

        last = para_rec._element
        for i, rec in enumerate(recs, 1):
            last = _add_para_after(last, f"  {i}.  {rec}", size=11, color=RGB_GRAY)

    # ── Glossaire ─────────────────────────────────────────────────────────────
    if para_glo:
        last = para_glo._element
        for term, definition in GLOSSAIRE:
            last = _add_para_after(last, f"* {term} :", bold=True, size=11, color=RGB_DARK)
            last = _add_para_after(last, f"   {definition}", size=11, color=RGB_GRAY)
            last = _add_para_after(last, "", size=5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_para_bg(p_elem, hex_color: str) -> None:
    """Set paragraph background via shading on paragraph properties."""
    # Trouve ou cree w:pPr
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    pPr = p_elem.find(f"{{{ns}}}pPr")
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_color)
    shd.set(qn("w:color"), hex_color)
    shd.set(qn("w:val"),   "clear")
    pPr.append(shd)
